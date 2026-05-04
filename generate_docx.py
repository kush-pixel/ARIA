"""Generate CS595_Project_Document_Final.docx from structured content."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def set_font(run, size=11, bold=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)


def add_heading(doc, text, level):
    para = doc.add_heading(text, level=level)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in para.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.bold = True
        if level == 1:
            run.font.size = Pt(14)
        elif level == 2:
            run.font.size = Pt(12)
        else:
            run.font.size = Pt(11)
    return para


def add_para(doc, text, bold=False, indent=0):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if indent:
        para.paragraph_format.left_indent = Inches(indent * 0.25)
    run = para.add_run(text)
    set_font(run, bold=bold)
    return para


def add_bullet(doc, text, level=0, bold_prefix=None):
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    if bold_prefix:
        r1 = para.add_run(bold_prefix)
        set_font(r1, bold=True)
    r2 = para.add_run(text)
    set_font(r2)
    return para


def add_labeled(doc, label, text, indent=0.25):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(indent)
    r1 = para.add_run(label + "  ")
    set_font(r1, bold=True)
    r2 = para.add_run(text)
    set_font(r2)
    return para


def build_doc():
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(0, 0, 0)

    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # TITLE BLOCK
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("CS595 2026 - Project Information Document")
    r.font.name = "Times New Roman"; r.font.size = Pt(16); r.font.bold = True
    r.font.color.rgb = RGBColor(0, 0, 0)

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run("ARIA - Adherence Risk Intelligence Agent")
    r.font.name = "Times New Roman"; r.font.size = Pt(14); r.font.bold = True
    r.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_paragraph()

    # ------------------------------------------------------------------
    # 1. PROJECT DEFINITION
    # ------------------------------------------------------------------
    add_heading(doc, "1. Project Definition", 1)
    add_labeled(doc, "Project Title:", "ARIA - Adherence Risk Intelligence Agent", indent=0)
    add_labeled(doc, "Author(s):", "Krishna Patel, Kush Patel, Sahilsingh Khalsa, Nesh Rochwani, Prakriti Sharma, Yash Sharma", indent=0)
    add_labeled(doc, "Team Name:", "NeuraCare Nexus", indent=0)
    para = doc.add_paragraph()
    r1 = para.add_run("LOF Pillar:  ")
    set_font(r1, bold=True)
    r2 = para.add_run("[ ] Patient Engagement     [X] Data Analysis & Population Health     [ ] Medical Education")
    set_font(r2)
    add_labeled(doc, "Last Updated:", "May 3, 2026", indent=0)
    doc.add_paragraph()

    # ------------------------------------------------------------------
    # 2. INTRODUCTION
    # ------------------------------------------------------------------
    add_heading(doc, "2. Introduction", 1)

    # 2a
    add_heading(doc, "2a. Problem Description", 2)
    add_para(doc, (
        "Clinicians managing large hypertension panels have no structured view of what happened between appointments. "
        "When a patient arrives, the clinician has only a single clinic BP reading, a static medication list, and no home monitoring record. "
        "Home readings, missed doses, and worsening trends remain invisible until the patient is physically present."
    ))
    add_para(doc, (
        "Blood pressure is controlled in fewer than half of treated hypertensive patients, with therapeutic inertia "
        "-- failure to adjust treatment when BP stays elevated -- as a key contributor. "
        "Without structured longitudinal data, clinicians cannot determine whether an elevated reading represents a "
        "sustained pattern warranting action or a transient spike."
    ))
    add_para(doc, (
        "Without adherence data linked to BP trajectory, clinicians cannot distinguish a patient whose BP is uncontrolled "
        "due to missed doses from one whose treatment is biologically insufficient -- two presentations demanding "
        "completely different clinical responses."
    ))
    add_para(doc, (
        "ARIA addresses these problems by providing a structured pre-visit clinical intelligence briefing before the "
        "patient arrives, so the clinician can act on the information rather than search for it."
    ))
    doc.add_paragraph()

    # 2b
    add_heading(doc, "2b. Purpose", 2)
    add_para(doc, (
        "ARIA is a between-visit clinical intelligence platform for hypertension management at general practice scale. "
        "It ingests structured EHR data, generates synthetic home monitoring records, analyzes the combined data through "
        "a three-layer AI pipeline, and delivers a structured pre-visit briefing organized by clinical priority."
    ))
    add_para(doc, "The project has the following specific goals and objectives:", bold=True)
    goals = [
        ("Detect reading gaps before they become clinical risks ",
         "using tier-appropriate thresholds (High: 1 day flag / 3 day urgent; Medium: 3 / 5; Low: 7 / 14)."),
        ("Identify therapeutic inertia using patient-specific thresholds ",
         "derived from the patient's own stable historic BP, with a -7 mmHg comorbidity adjustment for CHF/Stroke/TIA."),
        ("Surface adherence concerns with clinical nuance, ",
         "distinguishing Pattern A (high BP + <80% adherence), Pattern B (high BP + high adherence), and Pattern C (normal BP + low adherence)."),
        ("Detect BP deterioration ",
         "using three simultaneous gates (positive slope, recent 3-day avg vs days 4-10, recent avg vs adaptive threshold) plus a step-change sub-detector."),
        ("Flag drug interactions deterministically ",
         "using four rules at every briefing generation with comorbidity-escalated severity and no LLM involvement."),
        ("Compute a numeric risk priority score (0.0-100.0) ",
         "from five weighted clinical signals to control secondary sort order within each risk tier."),
        ("Generate a validated, readable briefing narrative ",
         "using a constrained LLM with eleven validation checks; on failure the full deterministic briefing is shown without an AI summary."),
        ("Validate system outputs against real physician clinical judgments ",
         "via shadow mode replay of Layer 1 detectors at historical clinic visits in the iEMR dataset."),
    ]
    for bold_part, rest in goals:
        add_bullet(doc, rest, bold_prefix=bold_part)
    doc.add_paragraph()

    # 2c
    add_heading(doc, "2c. Scope", 2)
    add_para(doc, "Included in Scope", bold=True)
    add_labeled(doc, "Disease Focus:", "Hypertension is the sole disease module; the architecture is extensible but no other disease has been implemented.", indent=0)
    add_labeled(doc, "EHR Integration via FHIR R4:", "Patient EHR data is ingested exclusively via FHIR R4 Bundle through an adapter layer that converts iEMR JSON, ensuring source-EHR changes require only a new adapter.", indent=0)
    add_labeled(doc, "Three-Layer AI Pipeline:", "Layer 1 is deterministic rule-based detection, Layer 2 computes a weighted numeric risk score, and Layer 3 uses an LLM to generate a readable narrative; layers run in strict sequence.", indent=0)
    pipeline_items = [
        ("Layer 1 - Deterministic Rule Engine: ", "Five detectors (gap, inertia, adherence, deterioration, variability) -- pure logic, no probabilistic components."),
        ("Layer 2 - Weighted Risk Scoring: ", "Five signals weighted to 0.0-100.0: systolic vs baseline (30%), medication inertia (25%), inverted adherence (20%), gap (15%), comorbidities (10%)."),
        ("Layer 3 - LLM Narrative: ", "Three-sentence summary generated after the deterministic briefing is stored; eleven checks must pass or readable_summary is stored as null."),
    ]
    for bp, rest in pipeline_items:
        add_bullet(doc, rest, level=1, bold_prefix=bp)
    add_labeled(doc, "Drug Interaction Detection:", "A deterministic rule checker evaluates four interaction rules at every briefing generation with no LLM involvement.", indent=0)
    add_labeled(doc, "Risk Tier System:", "Every patient carries a High/Medium/Low tier set at ingestion for CHF/Stroke/TIA and maintained by a nightly reclassification job with hysteresis; clinicians can override with a documented reason.", indent=0)
    add_labeled(doc, "Clinician Dashboard:", "A Next.js 14 web app with patient list sorted by tier/score, individual briefing pages, alert inbox with disposition capture, and an embedded clinical chatbot.", indent=0)
    add_labeled(doc, "Patient PWA - CuffLink:", "An installable PWA for home BP submission and medication dose confirmation, with a separate JWT secret from the clinician dashboard.", indent=0)
    add_labeled(doc, "Synthetic Data Engine:", "Clinically realistic home BP readings and medication confirmations generated by interpolating between clinic anchor points with Gaussian noise and clinical realism rules.", indent=0)
    add_labeled(doc, "Shadow Mode Validation:", "An offline pipeline replaying Layer 1 detectors at historical clinic visits and comparing output against physician PROBLEM_STATUS2_FLAG labels.", indent=0)
    add_labeled(doc, "Full Audit Logging:", "An immutable audit_events table records every significant system action with actor, patient, action, resource, and outcome ('success' or 'failure').", indent=0)
    add_labeled(doc, "Background Processing:", "A polling worker processes queued jobs every 30 seconds; APScheduler drives briefing generation at 7:30 AM and a midnight pattern recompute sweep.", indent=0)
    add_labeled(doc, "Alert Feedback Loop:", "Clinicians acknowledge alerts with a disposition (agree_acting/agree_monitoring/disagree); Disagree auto-schedules a 30-day outcome verification, and four or more Disagree dispositions surface a calibration recommendation.", indent=0)
    doc.add_paragraph()
    add_para(doc, "Excluded from Scope", bold=True)
    excluded = [
        "OpenMRS integration - identified as a stretch goal; not implemented.",
        "BLE Bluetooth cuff auto-submission - webhook endpoint exists but no device pairing implemented.",
        "Real-time pharmacy data feeds, electronic prescription access, or live dispensing records.",
        "Machine learning models for prediction, forecasting, or adaptive threshold calibration.",
        "Type 2 Diabetes extension module - defined in early planning documents but not implemented.",
        "Full production EHR integration beyond the iEMR prototype dataset.",
        "HIPAA-compliant production security hardening or enterprise deployment infrastructure.",
        "Multi-patient shadow mode validation - shadow mode validated on patient 1091 only.",
        "SMS submission pathway for feature phones.",
    ]
    for e in excluded:
        add_bullet(doc, e)
    doc.add_paragraph()

    # 2d
    add_heading(doc, "2d. Audience", 2)
    add_para(doc, "Primary Users - Clinicians (GPs, Internists, Cardiologists, Nurse Practitioners, Care Managers)", bold=True)
    add_para(doc, (
        "The primary users are clinicians who manage hypertension patient panels under significant time pressure, "
        "with no practical way to manually track between-visit history across hundreds of patients. "
        "Every ARIA finding includes the specific data that produced it, and clinicians can override risk tiers or disagree with any finding."
    ))
    doc.add_paragraph()
    add_para(doc, "Secondary Users - Population Health Managers and Care Coordinators", bold=True)
    add_para(doc, (
        "Population health managers use ARIA at the panel level to identify the highest-risk patient segments "
        "and systemic issues such as widespread therapeutic inertia across the panel."
    ))
    doc.add_paragraph()
    add_para(doc, "Tertiary Users - Health Informatics Researchers and Academic Reviewers", bold=True)
    add_para(doc, (
        "Academic reviewers are interested in the technical rigor of the AI pipeline, clinical validity of "
        "detection logic, FHIR R4 standards compliance, and shadow mode validation methodology."
    ))
    doc.add_paragraph()

    # 2e
    add_heading(doc, "2e. Overview", 2)
    add_para(doc, (
        "ARIA is a clinical intelligence platform built for the gap between EHR visits. "
        "It consists of five integrated components that transform raw EHR and home monitoring data into a structured, "
        "prioritized briefing delivered before the patient arrives."
    ))
    doc.add_paragraph()
    components = [
        ("EHR Integration - ", "The iEMR adapter converts raw JSON into FHIR R4 Bundle format; the ingestion layer populates 12 tables and triggers immovable risk tier overrides for CHF/Stroke/TIA patients."),
        ("Synthetic Data Engine - ", "Generates clinically realistic home BP readings across the patient's full care history by interpolating between clinic anchor points with Gaussian noise, morning/evening differential, device outage gaps, and a white-coat dip before each appointment."),
        ("Three-Layer AI Pipeline - ", "Layer 1 runs five deterministic detectors; Layer 2 computes a weighted risk score; Layer 3 generates a three-sentence LLM narrative subject to eleven validation checks."),
        ("Clinician Dashboard - ", "A Next.js 14 app at port 3000 with server-side paginated patient list (25/page), individual briefing pages, alert inbox with structured disposition capture, and the Ask ARIA clinical chatbot."),
        ("Patient PWA - CuffLink - ", "An installable Next.js 14 PWA at port 3001 for home BP submission and medication dose confirmation; patients see no clinical interpretation, risk scores, or readings."),
    ]
    for bp, rest in components:
        add_bullet(doc, rest, bold_prefix=bp)
    doc.add_paragraph()
    add_para(doc, (
        "All ARIA outputs are framed as decision support only; clinical language constraints (no 'non-adherent', "
        "no dose-change instructions) are enforced in code at both the deterministic and LLM validation layers."
    ))
    doc.add_paragraph()

    # 2f
    add_heading(doc, "2f. Glossary / Terminology", 2)
    glossary = [
        ("FHIR R4", "HL7 standard for healthcare data interchange; ARIA uses FHIR R4 Bundle as the adapter-to-ingestion interchange format."),
        ("iEMR", "De-identified longitudinal clinical record set from Leap of Faith Technologies containing clinic BP readings, medication history, physician HTN concern assessments, and per-visit problem assessments."),
        ("Therapeutic Inertia", "Failure to intensify or adjust antihypertensive treatment when BP remains persistently elevated above the patient's clinical target."),
        ("Layer 1 / Layer 2 / Layer 3", "Sequential AI pipeline stages: deterministic rule-based detection (Layer 1), weighted numeric risk scoring (Layer 2), and LLM narrative generation with eleven validation checks (Layer 3)."),
        ("Risk Tier", "Categorical patient classification (High/Medium/Low) that is the primary dashboard sort key, set at ingestion for qualifying conditions and maintained algorithmically."),
        ("Risk Score", "Numeric priority value 0.0-100.0 from Layer 2 controlling sort order within each tier; five signals: systolic vs baseline (30%), medication inertia (25%), inverted adherence (20%), gap (15%), comorbidities (10%)."),
        ("Adaptive Window", "Variable-length detection window (14-90 days) derived from last_visit_date to next_appointment, used by all five Layer 1 detectors; falls back to 28 days when appointment data is unavailable."),
        ("Patient-Adaptive Threshold", "Individual systolic threshold: max(130, stable_baseline_mean + 1.5 x SD) capped at 145, reduced by 7 mmHg for CHF/Stroke/TIA with a floor of 130."),
        ("White-Coat Dip", "Observed BP drop before clinical appointments; inertia and deterioration detectors exclude readings within 5 days of next_appointment to prevent masking elevated patterns."),
        ("Drug Interaction Detector", "Deterministic checker evaluating four rules (NSAID+antihypertensive, triple whammy, K-sparing+ACE/ARB, BB+non-DHP CCB) with comorbidity-escalated severity."),
        ("Shadow Mode", "Offline validation replaying Layer 1 detectors at each historical clinic visit using only pre-visit data, comparing output against physician PROBLEM_STATUS2_FLAG labels."),
        ("CuffLink", "The patient-facing PWA brand name for ARIA's home monitoring layer; currently supports manual entry with BLE device pairing on the roadmap."),
        ("Cold-Start Suppression", "Prevents inertia, deterioration, and adherence detectors from running for patients enrolled fewer than 21 days; gap detector is exempt."),
        ("Mini-Briefing", "Between-visit briefing triggered by an urgent alert, with null appointment_date, always treated as active by the briefing API."),
        ("Calibration Suppression", "When four or more Disagree dispositions are recorded for a patient-detector pair and an approved calibration_rules row exists, inbox alert writes for that pair are suppressed."),
    ]
    for term, definition in glossary:
        para = doc.add_paragraph()
        r1 = para.add_run(term + ": ")
        set_font(r1, bold=True)
        r2 = para.add_run(definition)
        set_font(r2)
    doc.add_paragraph()

    # 2g
    add_heading(doc, "2g. Skills Needed", 2)
    add_para(doc, "The project spans backend engineering, database design, frontend development, clinical informatics, LLM integration, and statistical signal processing.")
    skills = [
        "Python 3.11, FastAPI async, SQLAlchemy 2.0 async ORM (select() only), Pydantic v2, pytest fixture-based testing.",
        "PostgreSQL -- 12-table schema design, index strategy, asyncpg driver, Supabase hosting.",
        "TypeScript, Next.js 14, Tailwind CSS, recharts for BP charts, PWA configuration with @ducanh2912/next-pwa.",
        "FHIR R4 -- Bundle structure, resource types (Patient, Condition, MedicationRequest, Observation), LOINC coding.",
        "Clinical informatics -- hypertension thresholds (NICE NG136), therapeutic inertia, drug interactions, ICD-10, RxNORM.",
        "LLM API integration -- Anthropic SDK, prompt design, output validation, tool-call patterns, prompt injection safeguards.",
        "Background job processing -- APScheduler, polling worker pattern, idempotency key design.",
        "JWT authentication -- separate clinician and patient secrets, token expiry, role enforcement.",
        "Statistical signal processing -- rolling averages, linear slope via least-squares, coefficient of variation, Beta distribution.",
    ]
    for s in skills:
        add_bullet(doc, s)
    doc.add_paragraph()

    # ------------------------------------------------------------------
    # 3. SYSTEM OVERVIEW
    # ------------------------------------------------------------------
    add_heading(doc, "3. System Overview", 1)

    # 3a
    add_heading(doc, "3a. Architecture", 2)
    add_para(doc, (
        "ARIA is structured as loosely coupled services sharing a single PostgreSQL database. "
        "No analytical computation occurs in the HTTP request path; all patient-level analysis runs asynchronously through the background worker."
    ))
    doc.add_paragraph()
    arch_groups = [
        ("EHR Ingestion Layer", [
            "adapter.py converts iEMR JSON to FHIR R4 Bundle; ingestion.py populates all 12 tables.",
            "Ingestion-time auto-overrides set risk_tier='high' and tier_override_source='system' for CHF/Stroke/TIA.",
            "Two-step upsert handles re-ingestion safely; system conditions always promote.",
        ]),
        ("Background Analysis Pipeline", [
            "processing_jobs table is the job queue; polling worker (processor.py) picks up queued jobs every 30 seconds.",
            "pattern_recompute jobs run all five Layer 1 detectors, write alerts, compute Layer 2 risk score, then call _apply_tier_reclassification().",
            "off_hours flag is stamped from pattern-onset reading datetime, not job execution time.",
        ]),
        ("Scheduled Operations (APScheduler)", [
            "7:30 AM daily: briefing_generation jobs queued for monitoring_active patients with appointments that day.",
            "Midnight UTC: pattern_recompute jobs spread across 0-119 minutes via deterministic MD5 offset per patient_id.",
            "scheduler.py runs as a separate process from the polling worker.",
        ]),
        ("Briefing Generation", [
            "composer.py assembles the deterministic briefing JSON including drug_interactions, trend_avg_systolic, and adaptive-window findings.",
            "summarizer.py calls the LLM with no DB session held open; llm_validator.py runs eleven checks before storage.",
            "On validation failure after one retry, readable_summary is stored as null and the full deterministic briefing is displayed.",
        ]),
        ("Clinician Dashboard (Next.js 14, port 3000)", [
            "GET /api/patients supports server-side pagination (25/page), search, and tier filtering.",
            "POST /api/alerts/{id}/acknowledge accepts disposition and auto-schedules 30-day outcome verification on Disagree.",
            "POST /api/chat handles Ask ARIA queries with tool-call routing and three-layer guardrails.",
        ]),
        ("Patient PWA - CuffLink (Next.js 14, port 3001)", [
            "POST /api/readings stores home BP; POST /api/confirmations/confirm records dose confirmations.",
            "GET /api/confirmations/ics returns a downloadable .ics medication reminder calendar file.",
            "Patients authenticate with a separate patient JWT (8h expiry, role='patient').",
        ]),
    ]
    for group_title, lines in arch_groups:
        para = doc.add_paragraph()
        r = para.add_run(group_title)
        set_font(r, bold=True)
        for line in lines:
            add_bullet(doc, line, level=1)
        doc.add_paragraph()

    # 3b
    add_heading(doc, "3b. Technologies Used", 2)
    tech_groups = [
        ("Backend", [
            "Python 3.11 with FastAPI -- async HTTP framework; all routes use async/await with Depends() for session injection.",
            "SQLAlchemy 2.0 async -- all queries use select() with await session.execute(); session.query() is prohibited.",
            "Pydantic v2 -- settings via SettingsConfigDict; asyncpg as the native PostgreSQL async driver.",
            "APScheduler -- 7:30 AM briefing jobs and midnight pattern recompute sweeps.",
            "ruff -- linting and formatting; ruff check app/ must pass before any work is complete.",
        ]),
        ("Database", [
            "PostgreSQL via Supabase -- 12 tables with asyncpg driver; all DB I/O is async.",
            "Key indexes: unique idempotency (processing_jobs), unique reading dedup (patient_id, effective_datetime, source), composite risk sort (risk_tier, risk_score DESC).",
        ]),
        ("AI and LLM", [
            "Anthropic claude-sonnet-4-20250514 -- specification model for Layer 3 and the Ask ARIA chatbot.",
            "OpenAI gpt-4o-mini -- current demo substitute; reversion to Anthropic is a one-line change per file.",
        ]),
        ("Frontend - Clinician Dashboard", [
            "Next.js 14, TypeScript strict mode, Tailwind CSS, recharts for BP trend charts.",
            "All API calls through src/lib/api.ts; all shared types in src/lib/types.ts.",
        ]),
        ("Patient PWA - CuffLink", [
            "Next.js 14 with @ducanh2912/next-pwa -- service worker and web app manifest enabling installation on iOS, Android, and desktop.",
        ]),
        ("Testing", [
            "pytest -- 601 fixture-based unit tests; no real patient data used in any test.",
            "@pytest.mark.integration separates integration tests from the standard test run.",
        ]),
    ]
    for group, items in tech_groups:
        para = doc.add_paragraph()
        r = para.add_run(group + ":")
        set_font(r, bold=True)
        for item in items:
            add_bullet(doc, item, level=1)
        doc.add_paragraph()

    # 3c
    add_heading(doc, "3c. Dependencies", 2)
    add_para(doc, "External Service Dependencies:", bold=True)
    ext_deps = [
        ("Supabase: ", "Managed PostgreSQL hosting; DATABASE_URL configured in backend/.env."),
        ("Anthropic API: ", "Layer 3 narrative generation and chatbot; currently substituted by OpenAI in demo deployment."),
        ("iEMR Dataset (Leap of Faith Technologies): ", "De-identified longitudinal records at data/raw/iemr/ -- the sole source of real clinical data."),
    ]
    for bp, rest in ext_deps:
        add_bullet(doc, rest, bold_prefix=bp)
    doc.add_paragraph()
    add_para(doc, "Internal Ordering Dependencies:", bold=True)
    add_bullet(doc, "EHR ingestion must complete before synthetic data generation (generator reads historic_bp_systolic from clinical_context).")
    add_bullet(doc, "Layer 1 must complete before Layer 2; Layer 2 before tier reclassification; full deterministic briefing must be stored before Layer 3 LLM call.")
    doc.add_paragraph()

    # 3d
    add_heading(doc, "3d. Hardware and Software Requirements", 2)
    add_para(doc, "Supported Platforms:", bold=True)
    add_bullet(doc, "macOS, Windows (10+), or Linux with Python 3.11 and Node.js.")
    add_para(doc, "Required Software:", bold=True)
    add_bullet(doc, "Python 3.11 via Conda (aria environment), Node.js LTS, Git.")
    add_bullet(doc, "A PostgreSQL client (psql, TablePlus, or DBeaver) for verifying database state.")
    add_para(doc, "Runtime Processes:", bold=True)
    add_bullet(doc, "uvicorn app.main:app --reload --port 8000 (from backend/)")
    add_bullet(doc, "npm run dev (from frontend/, port 3000) and npm run dev (from patient-app/, port 3001)")
    add_bullet(doc, "python scripts/run_worker.py and python scripts/run_scheduler.py")
    add_para(doc, "Credentials (backend/.env):", bold=True)
    add_bullet(doc, "DATABASE_URL, ANTHROPIC_API_KEY (or OPENAI_API_KEY for demo), JWT_SECRET, PATIENT_JWT_SECRET.")
    doc.add_paragraph()

    # 3e
    add_heading(doc, "3e. Data Description and Sources", 2)
    add_para(doc, "Primary Source - iEMR Dataset:", bold=True)
    add_para(doc, (
        "The iEMR dataset is a de-identified longitudinal record set from Leap of Faith Technologies containing "
        "clinic BP readings, full medication history, encounter dates, physician HTN concern assessments "
        "(PROBLEM_STATUS2_FLAG: 1=urgent, 2=concerned, 3=stable), allergies, lab orders, and problem assessments. "
        "The primary demo patient (1091) has 124 visits over 11 years, 65 clinic BP readings, 14 medications "
        "including a triple whammy interaction, and 17 coded problems including CHF (I50.9)."
    ))
    doc.add_paragraph()
    add_para(doc, "Synthetic Home Reading and Confirmation Data:", bold=True)
    add_para(doc, (
        "Home BP readings are generated by interpolating linearly between consecutive clinic anchor points with "
        "Gaussian noise (SD 8-12 mmHg), morning/evening differential, device outage gaps as absent rows, and "
        "a white-coat dip before each appointment. "
        "Medication confirmation records use a Beta distribution (mean ~91%) with per-medication per-interval variation."
    ))
    doc.add_paragraph()
    add_para(doc, "Demo Patient Set:", bold=True)
    demo_pts = [
        ("Patient A (1091) - Therapeutic Inertia + Drug Interaction: ",
         "High Risk (CHF system override); triple whammy critical interaction (Voltaren + ramipril + furosemide); inertia alert from sustained elevation without med change in window."),
        ("Patient B (DEMO_EHR) - EHR-Only: ",
         "monitoring_active=False; NSAID+antihypertensive interaction; overdue lab flag; data limitations banner."),
        ("Patient C (DEMO_GAP / David Patel) - Reading Gap: ",
         "82 days of readings with last reading Apr 23; urgent gap alert (9+ days by May 2) and inertia alert."),
        ("Patient D (DEMO_ADH) - Adherence Concern: ",
         "~58% adherence over 28 days with ~152 mmHg avg; Pattern A adherence alert and inertia alert."),
    ]
    for bp, rest in demo_pts:
        add_bullet(doc, rest, bold_prefix=bp)
    doc.add_paragraph()

    # 3f
    add_heading(doc, "3f. Pre-Work / Setup", 2)
    steps = [
        ("Step 1 - Environment: ", "conda activate aria; pip install -r requirements.txt (from backend/); npm install (from frontend/ and patient-app/)."),
        ("Step 2 - Credentials: ", "Create backend/.env with DATABASE_URL, ANTHROPIC_API_KEY or OPENAI_API_KEY, JWT_SECRET, PATIENT_JWT_SECRET."),
        ("Step 3 - Database: ", "python scripts/setup_db.py -- creates all 12 tables and indexes (safe to re-run)."),
        ("Step 4 - EHR Ingestion: ", "python scripts/run_adapter.py --patient 1091 then python scripts/run_ingestion.py --patient 1091."),
        ("Step 5 - Synthetic Data: ", "python scripts/run_generator.py --patient 1091 --mode full-timeline."),
        ("Step 6 - Demo Setup: ", "python scripts/setup_demo.py -- seeds all four demo patients; success = ALL CHECKS PASSED for all four."),
        ("Step 7 - Start Services: ", "uvicorn (8000), npm dev frontend (3000), npm dev patient-app (3001), run_worker.py, run_scheduler.py."),
    ]
    for bp, rest in steps:
        add_bullet(doc, rest, bold_prefix=bp)
    doc.add_paragraph()

    # ------------------------------------------------------------------
    # 4. FUNCTIONAL REQUIREMENTS
    # ------------------------------------------------------------------
    add_heading(doc, "4. Functional Requirements", 1)

    # 4a
    add_heading(doc, "4a. List of User Tasks", 2)
    task_groups = [
        ("Clinician Tasks", [
            "View patient panel sorted by clinical urgency (tier then risk score)",
            "Filter patient list by risk tier (All / High / Medium / Low)",
            "Search for a patient by name or patient ID",
            "Open and read an individual patient pre-visit briefing",
            "Review drug interaction flags, medication status, and adherence signal",
            "Review active problems, overdue investigations, and prioritized visit agenda",
            "Acknowledge an alert with a structured disposition (Agree: acting / Agree: monitoring / Disagree)",
            "Override a patient's risk tier with a documented reason",
            "Ask a natural-language clinical question via the Ask ARIA chatbot",
            "Review calibration recommendations and approve calibration rules",
            "Respond to 30-day outcome verification prompts",
        ]),
        ("Patient Tasks", [
            "Log in to the CuffLink patient app",
            "Submit a home BP reading with session type and optional symptoms",
            "Confirm a scheduled medication dose with a single tap",
            "Download a medication schedule calendar reminder file (.ics)",
        ]),
        ("Admin / System Tasks", [
            "Ingest a FHIR R4 Bundle via POST /api/ingest",
            "Trigger manual briefing generation and pattern recompute via the admin interface",
            "Run shadow mode evaluation via python scripts/run_shadow_mode.py",
        ]),
    ]
    for group, tasks in task_groups:
        para = doc.add_paragraph()
        r = para.add_run(group + ":")
        set_font(r, bold=True)
        for t in tasks:
            add_bullet(doc, t, level=1)
        doc.add_paragraph()

    # 4b
    add_heading(doc, "4b. User Stories", 2)
    stories = [
        ("As a clinician preparing for morning appointments, ",
         "I want to see which patients need the most attention today, sorted by clinical urgency."),
        ("As a clinician reviewing a briefing before a consultation, ",
         "I want to see drug interaction flags with plain-English descriptions so I can address them without a separate lookup."),
        ("As a clinician seeing a patient with persistently elevated BP, ",
         "I want to know whether the patient has been taking their medications consistently so I can distinguish an adherence issue from treatment failure."),
        ("As a clinician who disagrees with the system's risk tier for a patient, ",
         "I want to demote the tier with a documented reason and have that decision respected for a defined period."),
        ("As a clinician reviewing a complex patient, ",
         "I want to ask the system a natural-language question and receive a direct, data-grounded answer."),
        ("As a clinician with limited time per consultation, ",
         "I want a prioritized visit agenda showing the three to six most important items to discuss."),
        ("As a patient managing hypertension at home, ",
         "I want to submit my morning and evening BP readings and confirm my medications in under a minute."),
        ("As a patient who sometimes forgets my medication schedule, ",
         "I want to download a calendar reminder file for my phone."),
    ]
    for bp, rest in stories:
        add_bullet(doc, rest, bold_prefix=bp)
    doc.add_paragraph()

    # 4c
    add_heading(doc, "4c. Use Cases", 2)
    use_cases = [
        ("Use Case 1: Clinician Morning Dashboard Review", [
            ("Actor:", "Clinician"),
            ("Pre-condition:", "At least one monitoring-active patient has a briefing generated for today's appointment."),
            ("Main Flow:", "Clinician opens localhost:3000. Patient list loads sorted by tier then score. Clinician opens highest-priority briefing, reviews AI summary, BP trend, drug interactions, and visit agenda."),
            ("Post-condition:", "briefings.read_at is set; audit_events row written with action='briefing_viewed'."),
            ("Alternative:", "If LLM validation failed, the AI summary card is hidden; full deterministic briefing is shown."),
        ]),
        ("Use Case 2: Clinician Alert Acknowledgment with Disposition", [
            ("Actor:", "Clinician"),
            ("Pre-condition:", "An active unacknowledged alert is visible in the inbox."),
            ("Main Flow:", "Clinician reviews the alert, selects a disposition (Agree: acting / Agree: monitoring / Disagree), optionally enters a reason, and clicks Acknowledge."),
            ("Post-condition:", "acknowledged_at set; disposition stored in alert_feedback; audit logged."),
            ("Disagree path:", "outcome_tracker.py creates an outcome_verifications row with check_after = acknowledged_at + 30 days."),
        ]),
        ("Use Case 3: Clinician Risk Tier Override", [
            ("Actor:", "Clinician (authenticated with clinician JWT)"),
            ("Pre-condition:", "Patient does not have a system-override tier (no CHF/Stroke/TIA)."),
            ("Main Flow:", "Clinician submits PATCH /api/patients/{id}/tier with tier value and reason. Demotion sets tier_override_suppressed_until = now + 28 days."),
            ("Post-condition:", "Tier updated; suppression set for demotions; nightly reclassification respects 28-day window."),
            ("Exception:", "409 Conflict if tier_override_source='system'; 422 if tier value invalid or reason missing."),
        ]),
        ("Use Case 4: Patient Blood Pressure Submission", [
            ("Actor:", "Patient (authenticated via CuffLink PWA)"),
            ("Pre-condition:", "Patient holds a valid patient JWT (8h expiry, role='patient')."),
            ("Main Flow:", "Patient opens /submit, enters readings, selects session type, optionally records symptoms and medication status, submits."),
            ("Post-condition:", "Reading stored (source='manual'); audit_events row written (action='reading_ingested')."),
            ("Validation:", "Systolic 60-250 and diastolic 40-150 enforced; out-of-range rejected with 422."),
        ]),
        ("Use Case 5: Ask ARIA Chatbot Query", [
            ("Actor:", "Clinician (authenticated)"),
            ("Pre-condition:", "Clinician has a patient briefing page open."),
            ("Main Flow:", "Clinician types a question. Off-topic check (Layer 1) and scope check (Layer 2) run first. Chatbot executes tool calls (get_briefing/get_readings/get_alerts) and composes a data-grounded response validated by Layer 3."),
            ("Limitation:", "Scoped to one patient at a time; cross-patient queries are blocked."),
        ]),
    ]
    for uc_title, uc_fields in use_cases:
        para = doc.add_paragraph()
        r = para.add_run(uc_title)
        set_font(r, bold=True)
        for label, value in uc_fields:
            add_labeled(doc, label, value, indent=0.5)
        doc.add_paragraph()

    # 4d
    add_heading(doc, "4d. Prioritization", 2)
    add_para(doc, "Priority 1 - Core (built and working):", bold=True)
    add_para(doc, (
        "All five Layer 1 detectors, risk tier system with ingestion-time auto-overrides and nightly reclassification, "
        "Layer 2 weighted risk scoring, briefing composition and storage, drug interaction detector, Layer 3 LLM narrative "
        "with eleven validation checks, clinician dashboard, alert inbox, patient PWA, full audit logging, "
        "background worker and APScheduler, 601 passing unit tests."
    ))
    doc.add_paragraph()
    add_para(doc, "Priority 2 - Enhanced Features (built and working):", bold=True)
    add_para(doc, (
        "Ask ARIA clinical chatbot with three-layer guardrails, synthetic data engine with full care timeline, "
        "calibration suppression for repeated clinician dismissals, clinician tier override with 28-day suppression, "
        "alert escalation, off-hours tagging, medication .ics calendar files, interactive dashboard onboarding tour, "
        "shadow mode validation pipeline, trend_avg_systolic as single source of truth for BP Trend column."
    ))
    doc.add_paragraph()
    add_para(doc, "Priority 3 - Future Work (not built):", bold=True)
    add_para(doc, (
        "BLE Bluetooth cuff pairing, vendor cloud webhook integration, SMART on FHIR launch, TOTP MFA, "
        "empirically calibrated risk score weights, multi-patient shadow mode validation, structured physician-AI "
        "comparison UI, SMS submission for feature phones."
    ))
    doc.add_paragraph()

    # 4e
    add_heading(doc, "4e. List of Functional Requirements", 2)
    frs = [
        ("FR-01", "The system shall ingest patient EHR data exclusively via FHIR R4 Bundle; the adapter layer converts iEMR JSON so the ingestion layer never references source-EHR field names directly."),
        ("FR-02", "ICD-10 codes CHF (I50), Haemorrhagic Stroke (I61), Ischaemic Stroke (I63/I64), and TIA (G45) trigger immovable High Risk overrides at ingestion (tier_override_source='system'); the PATCH endpoint returns 409 for system-override patients."),
        ("FR-03", "Five Layer 1 detectors (gap, inertia, adherence, deterioration, variability) run nightly for every monitoring-active patient using the adaptive window and patient-adaptive threshold."),
        ("FR-04", "Gap thresholds by tier -- High: flag 1 day / urgent 3 days; Medium: 3 / 5; Low: 7 / 14. Gap_urgent alert generated when urgent threshold is exceeded; gap_briefing alert when flag threshold only is exceeded."),
        ("FR-05", "Inertia fires only when all four gates pass: rolling avg above patient threshold, at least five readings above threshold, elevated pattern spans more than 7 days, and no medication change from med_history JSONB falls within the adaptive window. White-coat exclusion removes readings within 5 days of next appointment."),
        ("FR-05a", "Variability detector computes coefficient of variation (CV = SD / mean x 100); CV >= 15% = high, >= 12% = moderate. Requires at least 7 readings; result appears in visit agenda with no separate alert row."),
        ("FR-06", "Adherence detector classifies Pattern A (high BP + adherence <80% -> 'possible adherence concern'), Pattern B (high BP + high adherence -> 'treatment review warranted', suppressed during active titration), or Pattern C (normal BP + low adherence -> 'contextual review')."),
        ("FR-07", "Deterioration fires when positive slope >= 0.3 mmHg/day, recent 3-day avg exceeds days 4-10 avg, and recent avg is at or above the adaptive threshold. A step-change sub-detector fires independently when the 7-day mean this week exceeds the 7-day mean from 3 weeks ago by >= 15 mmHg."),
        ("FR-08", "Layer 2 computes a 0.0-100.0 risk score: systolic vs personal baseline (30%), days since med change normalized to 180 days (25%), inverted adherence (20%), gap normalized to adaptive window (15%), severity-weighted comorbidities clamped to 0-100 (10%)."),
        ("FR-09", "risk_score and risk_score_computed_at are updated on the patients table after every Layer 2 computation; the dashboard shows a 'Score outdated' badge when risk_score_computed_at is more than 26 hours old."),
        ("FR-10", "Four drug interaction rules are evaluated at every briefing: NSAID+antihypertensive (warning/concern), triple whammy (concern/critical), K-sparing+ACE/ARB (warning/concern), BB+non-DHP CCB (always concern). Triple whammy deduplicates the simpler NSAID rule."),
        ("FR-11", "Layer 3 generates a three-sentence narrative; eleven checks must pass before storage. On retry failure, readable_summary is stored as null and every validation attempt writes an audit_events row with action='llm_validation'."),
        ("FR-12", "Layer 3 output is rejected for any of: 'non-adherent', 'non-compliant', 'hypertensive crisis', 'medication failure', dose-change instructions, 'prescribe', 'diagnose', 'emergency', 'tell the patient', patient ID verbatim, or prompt injection markers."),
        ("FR-13", "GET /api/briefings/{patient_id} returns only the most-recent active briefing (appointment_date >= today OR IS NULL); past appointment briefings are excluded and mini-briefings with null appointment_date are always active."),
        ("FR-14", "Patient list is sorted backend-side by risk_tier (High > Medium > Low) then risk_score DESC within tier, returned pre-sorted from GET /api/patients."),
        ("FR-15", "BP Trend in the dashboard uses trend_avg_systolic from the active briefing as the single source of truth; falls back to a live 28-day home-readings-only computation when null."),
        ("FR-16", "The patient PWA never shows readings, risk tier, risk score, or clinical interpretation; patient experience is limited to BP submission and medication confirmation."),
        ("FR-17", "Every bundle import, reading ingestion, briefing view, alert acknowledgment, LLM validation, and tier reclassification produces an audit_events row with outcome exactly 'success' or 'failure'."),
        ("FR-18", "Nightly reclassification uses hysteresis: medium to high at score >= 75; high to medium at score < 40 (system_score source only); medium to low requires score < 25 plus enrollment >= 90 days, no SEVERE/MODERATE comorbidity, and no active urgent alerts."),
        ("FR-19", "Clinician tier demotion sets tier_override_suppressed_until = now + 28 days (NICE NG136 section 1.6.3); nightly reclassification skips the patient during suppression unless score >= 85 (break-glass promotion only)."),
        ("FR-20", "The Ask ARIA chatbot applies three sequential guardrails: off-topic detection (Layer 1), scope check blocking cross-patient queries (Layer 2), and output validation blocking prescriptive language (Layer 3)."),
        ("FR-21", "Synthetic readings are generated with day-to-day SD 8-12 mmHg, morning readings 5-10 mmHg above evening, two-reading sessions where reading 2 is 2-6 mmHg lower, and device outages as absent rows with no null values."),
        ("FR-22", "Cold-start suppression prevents inertia, deterioration, and adherence detectors from running for patients enrolled fewer than 21 days; gap detector is exempt and runs from enrollment day one."),
        ("FR-23", "POST /api/alerts/{id}/acknowledge accepts an optional disposition (agree_acting/agree_monitoring/disagree) stored in alert_feedback. Four or more Disagree dispositions for a patient-detector pair surface a calibration recommendation via GET /api/admin/calibration-recommendations."),
        ("FR-24", "A Disagree disposition auto-creates an outcome_verifications row with check_after = acknowledged_at + 30 days; after 30 days the clinician responds via POST /api/admin/outcome-verifications/{id}/respond with 'relevant', 'not_relevant', or 'unsure'."),
    ]
    for fr_id, fr_text in frs:
        para = doc.add_paragraph()
        r1 = para.add_run(fr_id + ": ")
        set_font(r1, bold=True)
        r2 = para.add_run(fr_text)
        set_font(r2)
        doc.add_paragraph()

    # 4f
    add_heading(doc, "4f. Data Structures and Interfaces", 2)
    add_para(doc, "The system uses 12 PostgreSQL tables.")
    doc.add_paragraph()
    tables = [
        ("patients", "Central registry storing demographics, risk_tier, tier_override, tier_override_source, tier_override_suppressed_until, risk_score, monitoring_active, and scheduling fields; risk_tier and risk_score are the dashboard's primary and secondary sort keys."),
        ("clinical_context", "Pre-computed EHR context (one row per patient) including active_problems, current_medications, med_history JSONB (full medication timeline), allergies, historic BP arrays, overdue_labs, social_context, problem_assessments JSONB, and recent_labs JSONB."),
        ("readings", "Home and clinic BP readings with two reading pairs per session, computed averages, session (morning/evening/ad_hoc), source (generated/manual/ble_auto/clinic), symptoms TEXT[], and a unique index on (patient_id, effective_datetime, source)."),
        ("medication_confirmations", "Per-dose adherence records; confirmed_at is null for missed doses. Unique index on (patient_id, medication_name, scheduled_time) prevents duplicate generation."),
        ("alerts", "Detector findings with alert_type, gap_days, systolic_avg, triggered_at, acknowledged_at, off_hours flag, and escalated flag; gap_urgent and deterioration alerts unacknowledged more than 24 hours are promoted to escalated=True."),
        ("alert_feedback", "Clinician dispositions (agree_acting/agree_monitoring/disagree) on acknowledged alerts, with reason_text and clinician_id; used for calibration suppression."),
        ("briefings", "Structured pre-visit briefing payloads; llm_response JSONB holds the full deterministic briefing including drug_interactions, trend_avg_systolic, visit_agenda, and urgent_flags. Also stores model_version, prompt_hash, and read_at."),
        ("processing_jobs", "Job queue with idempotency_key UNIQUE, status lifecycle (queued to running to succeeded/failed), and job_type (pattern_recompute/briefing_generation/bundle_import)."),
        ("audit_events", "Immutable log with actor_type, actor_id, patient_id, action, resource_type, resource_id, and outcome (exactly 'success' or 'failure')."),
        ("gap_explanations", "Clinician-submitted gap explanations with reason (device_issue/travel/illness/unknown/non_compliance) and gap date range."),
        ("calibration_rules", "Approved suppression rules per patient-detector pair; when active, suppress inbox alert writes while detection still runs and findings still appear in the briefing."),
        ("outcome_verifications", "30-day retrospective checks auto-created on Disagree dispositions; clinician responds with 'relevant', 'not_relevant', or 'unsure' after check_after date."),
    ]
    for tname, tdesc in tables:
        para = doc.add_paragraph()
        r1 = para.add_run(tname + ": ")
        set_font(r1, bold=True)
        r2 = para.add_run(tdesc)
        set_font(r2)
        doc.add_paragraph()

    add_para(doc, "Key REST API Endpoints:", bold=True)
    endpoints = [
        "GET /api/patients - paginated (page_size=25), sorted by tier then score, server-side search and tier filter, includes trend_avg_systolic.",
        "GET /api/patients/{id} / PATCH /api/patients/{id}/tier - single patient; tier override with 28-day suppression and 409 for system overrides.",
        "GET /api/readings / POST /api/readings - reading history and home BP submission.",
        "GET /api/briefings/{patient_id} - most-recent active briefing (appointment_date >= today OR IS NULL).",
        "GET /api/alerts / GET /api/alerts/acknowledged - active and acknowledged alerts with patient names.",
        "POST /api/alerts/{id}/acknowledge / DELETE /api/alerts/{id}/acknowledge - acknowledge with disposition; undo within 24 hours.",
        "POST /api/ingest - FHIR Bundle ingestion.",
        "GET /api/confirmations/pending / POST /api/confirmations/confirm / GET /api/confirmations/ics - patient dose management.",
        "POST /api/chat / POST /api/chat/summary/{patient_id} / POST /api/chat/feedback - Ask ARIA chatbot.",
        "POST /api/admin/trigger-scheduler - manual briefing and recompute trigger.",
        "GET /api/admin/calibration-recommendations / POST /api/admin/calibration-rules - calibration workflow.",
        "GET /api/admin/outcome-verifications / POST /api/admin/outcome-verifications/{id}/respond - 30-day retrospective checks.",
        "GET /api/shadow-mode/{patient_id} - shadow mode evaluation results.",
    ]
    for ep in endpoints:
        add_bullet(doc, ep)
    doc.add_paragraph()

    # ------------------------------------------------------------------
    # 5. GENERAL SCHEDULE
    # ------------------------------------------------------------------
    add_heading(doc, "5. General Schedule", 1)
    add_para(doc, (
        "The project was developed across two sprints from January to May 2026 by six team members, "
        "with demo preparation completing on May 3, 2026."
    ))
    doc.add_paragraph()

    add_heading(doc, "Sprint 1 - January 2026 to March 2026", 2)
    s1 = [
        "Database schema design: all 12 tables, ADD COLUMN IF NOT EXISTS migration system, and all performance and idempotency indexes.",
        "EHR ingestion pipeline: iEMR adapter converting source JSON to FHIR R4 Bundle, ingestion.py populating all tables, ingestion-time risk tier auto-overrides for CHF/Stroke/TIA.",
        "Layer 1 detectors: gap, therapeutic inertia (patient-adaptive threshold, white-coat exclusion), adherence-BP correlation, and deterioration; adaptive window logic across all detectors.",
        "Layer 2 risk scoring: five-signal weighted formula in risk_scorer.py, severity-weighted comorbidity scoring, score stored on patients table.",
        "Background worker and basic clinician dashboard with patient list sorted by tier and score; initial shadow mode evaluation script.",
    ]
    for item in s1:
        add_bullet(doc, item)
    doc.add_paragraph()

    add_heading(doc, "Sprint 2 - March 2026 to May 2026", 2)
    s2 = [
        "Layer 3 LLM narrative (eleven validation checks, retry, prompt_hash per briefing), drug interaction detector (four rules, comorbidity escalation), and variability detector.",
        "Risk tier reclassification with hysteresis and break-glass, clinician tier override endpoint, alert escalation and off-hours tagging.",
        "Full briefing page UI, alert inbox with disposition capture, and Ask ARIA chatbot with three-layer guardrails and tool-call routing.",
        "Patient PWA CuffLink: login, BP submission with symptom logging, medication confirmation, .ics calendar download.",
        "All four demo patients seeded via setup_demo.py (ALL CHECKS PASSED); 601 unit tests passing; shadow mode at 78.4% agreement across 37 labeled evaluation points.",
    ]
    for item in s2:
        add_bullet(doc, item)
    doc.add_paragraph()

    add_heading(doc, "Final Project Completion Status", 2)
    add_para(doc, (
        "The project is complete for prototype purposes as of May 3, 2026. "
        "All five Layer 1 detectors, Layer 2 scoring, Layer 3 LLM validation, drug interaction detection, patient PWA, "
        "Ask ARIA chatbot, and all four demo patients are fully operational. "
        "Shadow mode reached 78.4% agreement (target 80%); the 1.6 pp gap reflects the limited evaluation dataset, "
        "with all six false negatives clinically explained and zero caused by system logic errors. "
        "Layer 3 and the chatbot currently use OpenAI gpt-4o-mini as a demo substitute; "
        "reversion to Anthropic claude-sonnet-4-20250514 is a one-line change per file."
    ))
    doc.add_paragraph()

    # ------------------------------------------------------------------
    # 6. SUBMISSION OF RESULTS
    # ------------------------------------------------------------------
    add_heading(doc, "6. Submission of Results", 1)

    add_heading(doc, "6a. Source Code", 2)
    add_para(doc, "All source code is maintained in the project Git repository (nesh branch).")
    src = [
        "backend/app/ -- FastAPI app with 12 ORM models, API routes, FHIR ingestion, pattern engine (5 detectors + risk scorer), briefing composer/validator/summarizer, chatbot agent, background worker, and scheduler.",
        "frontend/src/ -- Next.js 14 clinician dashboard: patient list, individual briefing page with SparklineChart, alert inbox, chatbot; all API calls through api.ts, all types in types.ts.",
        "patient-app/src/ -- CuffLink PWA: login, BP submit, medication confirm; @ducanh2912/next-pwa for installability on iOS, Android, and desktop.",
        "scripts/ -- run_adapter.py, run_ingestion.py, run_generator.py, setup_db.py, setup_demo.py, run_worker.py, run_scheduler.py, run_shadow_mode.py. tests/ -- 601 fixture-based unit tests. prompts/ -- briefing_summary_prompt.md, chat_system_prompt.md.",
    ]
    for item in src:
        add_bullet(doc, item)
    doc.add_paragraph()

    add_heading(doc, "6b. Data and Sample Data", 2)
    data_items = [
        "data/raw/iemr/ -- de-identified iEMR source data from Leap of Faith Technologies.",
        "All four demo patients seeded by python scripts/setup_demo.py (idempotent; success = ALL CHECKS PASSED).",
        "data/shadow_mode_results.json -- most recent shadow mode evaluation results for patient 1091.",
        "Synthetic readings and confirmations generated on demand via run_generator.py --mode full-timeline.",
    ]
    for item in data_items:
        add_bullet(doc, item)
    doc.add_paragraph()

    add_heading(doc, "6c. AI Prompts", 2)
    prompts = [
        ("prompts/briefing_summary_prompt.md - ", "Layer 3 system prompt specifying the three-sentence output structure, all forbidden phrases, faithfulness requirements, and the instruction set mirroring the eleven validator checks. SHA-256 hash stored in briefings.prompt_hash per generated summary."),
        ("prompts/chat_system_prompt.md - ", "Ask ARIA chatbot system prompt specifying the three-layer guardrail structure, tool schemas (get_briefing, get_readings, get_alerts), and single-patient scope boundary."),
        ("backend/app/services/chat/agent.py - ", "Chatbot implementation loading the system prompt from prompts/, routing tool calls, and enforcing guardrails."),
    ]
    for bp, rest in prompts:
        add_bullet(doc, rest, bold_prefix=bp)
    doc.add_paragraph()

    add_heading(doc, "6d. Documentation", 2)
    doc_items = [
        ("ARIA_MASTER_DOCUMENT.md - ", "Primary technical documentation covering the three-layer AI architecture, detection logic, risk tier system, shadow mode methodology, and future roadmap."),
        ("AUDIT.md - ", "Full log of over 60 architectural decisions and fixes with root cause analysis and shadow mode false positive/negative documentation."),
        ("STATUS.md - ", "Chronological development log recording each significant change, files modified, and outcomes."),
        ("documentation/ARIA_Full_Specification_v5_0.md - ", "Full system specification covering API contracts, data models, and clinical logic at production-readiness level."),
        ("documentation/ARIA_SOP_v2_0.md - ", "Standard operating procedure covering setup, worker operation, alert management, calibration, and troubleshooting."),
        ("documentation/ - ", "Supplementary docs: team_access_guide.md, patientapp.md, chatbot.md, risk_tier_reclassification_v1_0.md."),
    ]
    for bp, rest in doc_items:
        add_bullet(doc, rest, bold_prefix=bp)
    doc.add_paragraph()

    add_heading(doc, "6e. Screenshots and Demo Materials", 2)
    add_para(doc, (
        "The demo is structured around four distinct clinical scenarios, one per demo patient. "
        "All four patients are seeded via setup_demo.py and produce reproducible findings on demo day."
    ))
    doc.add_paragraph()
    add_para(doc, "Each demo patient demonstrates a distinct clinical scenario:", bold=True)
    demos = [
        ("Patient 1091 (High Risk - Therapeutic Inertia + Drug Interaction): ", "Shows the full briefing with a critical triple whammy drug interaction (Voltaren + ramipril + furosemide escalated by CHF+CKD), therapeutic inertia alert, CHF system-override tier explanation via tooltip, and Ask ARIA chatbot responding to drug interaction queries."),
        ("DEMO_EHR (EHR-Only Patient): ", "Shows the briefing for a monitoring_active=False patient demonstrating NSAID+antihypertensive interaction, overdue lab flag, and the data limitations banner."),
        ("DEMO_GAP / David Patel (Reading Gap): ", "Shows urgent gap alert (9-12 days by demo date), inertia alert from the preceding elevated window, and escalation badge if unacknowledged more than 24 hours."),
        ("DEMO_ADH (Adherence Concern): ", "Shows Pattern A adherence alert (~58% adherence with elevated BP), inertia alert, and per-medication confirmation rates on the briefing page."),
    ]
    for bp, rest in demos:
        add_bullet(doc, rest, bold_prefix=bp)
    doc.add_paragraph()

    add_heading(doc, "6f. Other Final Deliverables", 2)
    add_para(doc, "Shadow Mode Validation Report:", bold=True)
    add_para(doc, (
        "Results at data/shadow_mode_results.json: 78.4% agreement across 37 labeled evaluation points, "
        "6 false negatives (all clinically explained: cold-start suppression, active treatment response, same-day med changes), "
        "and 2 false positives documented in AUDIT.md. The 1.6 pp gap below the 80% target reflects the limited evaluation dataset."
    ))
    doc.add_paragraph()
    add_para(doc, "Architecture Decision Record:", bold=True)
    add_para(doc, (
        "AUDIT.md documents over 60 architectural decisions including adaptive window bounds, patient-adaptive threshold derivation, "
        "hysteresis band design (40-75), LLM validation regex choices, two-step upsert at ingestion, and the trend_avg_systolic "
        "briefing field replacing live DB computation to eliminate up to 14 mmHg dashboard divergence."
    ))
    doc.add_paragraph()
    add_para(doc, "Clinical Boundary Enforcement Record:", bold=True)
    add_para(doc, (
        "Clinical language constraints are enforced at the deterministic level (briefing composer strings) and at the LLM "
        "validation level (llm_validator.py eleven checks). The 409 response from PATCH /api/patients/{id}/tier enforces "
        "the boundary at the API level for system-override patients."
    ))
    doc.add_paragraph()
    add_para(doc, "Limitations Summary:", bold=True)
    limitations = [
        "Shadow mode at 78.4%, 1.6 pp below the 80% target; expanding the evaluation dataset would likely improve this.",
        "Layer 3 and chatbot use OpenAI gpt-4o-mini rather than Anthropic claude-sonnet-4-20250514; reversion is a one-line change.",
        "Full pipeline validated on a single de-identified patient (1091); multi-patient validation required before any clinical deployment.",
        "Medication confirmation data is synthetic; real deployment requires smart dispenser or pharmacist dispensing records.",
        "The five risk score weights are expert-informed, not empirically calibrated; calibration_rules infrastructure exists for future recalibration.",
        "The FHIR adapter is written specifically for the LOF iEMR JSON structure; a different EHR requires a new adapter.",
    ]
    for lim in limitations:
        add_bullet(doc, lim)

    out_path = "/Users/neshrochwani/Desktop/ARIA_CS595/ARIA/CS595_Project_Document_Final.docx"
    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    build_doc()
